import sys
import click
from rich.console import Console
from rich.table import Table
from rich.panel import Panel
from rich.text import Text
from rich.columns import Columns
from prompt_toolkit import PromptSession
from prompt_toolkit.formatted_text import HTML
from prompt_toolkit.styles import Style
from prompt_toolkit.history import FileHistory
from prompt_toolkit.completion import Completer, Completion
from jk_core.prompt_engine import assemble_prompt, get_required_vars
from jk_core.ai_client import GeminiClient
from jk_core.orchestrator import Orchestrator
from jk_core.session_manager import SessionManager
from jk_core.search_engine import SearchEngine

console = Console()

SLASH_COMMANDS = [
    "/branch", "/copy", "/edit", "/help", "/history", "/latest",
    "/export", "/model", "/name", "/new", "/proj", "/reset",
    "/resume", "/retry", "/save", "/search", "/stats", "/switch",
    "/temp", "/undo",
]

class _SlashCompleter(Completer):
    """
    Tab-completes slash commands at the start of input.
    """
    def get_completions(self, document, complete_event):
        text = document.text_before_cursor

        # Slash command completion at start of input
        if text.startswith("/"):
            for cmd in SLASH_COMMANDS:
                if cmd.startswith(text):
                    yield Completion(cmd, start_position=-len(text))

# Shared prompt_toolkit session — persists input history across runs
from jk_core.constants import SHARED_CONFIG_PATH
from pathlib import Path
_pt_history_file = Path(SHARED_CONFIG_PATH) / "prompt_history"
_pt_session: PromptSession = PromptSession(history=FileHistory(str(_pt_history_file)), completer=_SlashCompleter())

# Cyan bold to match existing rich palette
_pt_style = Style.from_dict({"prompt": "ansicyan bold"})

class ChatRouter:
    """
    Routes Slash Commands to specific logic.
    Each method handles a specific user intent.
    The cmd_<name> is ordered by name because it's easier to search that way.
    """
    def __init__(self, client, orchestrator, session_manager):
        self.client = client
        self.orch = orchestrator
        self.session = session_manager
        self.search_engine = SearchEngine(client)
        self._session_map = {}
        self._edit_map = {} # CHANGE: Initialize the map here

    def handle(self, text: str) -> bool:
        """
        Returns:
            - True: If a standard command was handled.
            - dict: If a Replay is triggered (contains future prompts).
            - False: If it's a regular message.
        """
        clean_text = text.strip().lower()

        exit_keywords = ["exit", "quit", "bye", "/exit", "/bye", "/quit"]
        if clean_text in exit_keywords:
            self.cmd_exit(None)

        if not text.startswith("/"):
            return False

        parts = text.split()
        cmd = parts[0].lower()
        args = parts[1:]

        # Dispatch Table
        commands = {
            "/branch": self.cmd_branch,
            "/help": self.cmd_help,
            "/export": self.cmd_export,
            "/history": self.cmd_history,
            "/copy": self.cmd_copy,
            "/edit": self.cmd_edit,
            "/latest": self.cmd_toggle_latest,
            "/model": self.cmd_model_info,
            "/name": self.cmd_name,
            "/new": self.cmd_new_chat,
            "/proj": self.cmd_switch_proj,
            "/reset": self.cmd_reset,
            "/resume": self.cmd_resume,
            "/retry": self.cmd_retry,
            "/save": self.cmd_save,
            "/search": self.cmd_search,
            "/stats": self.cmd_stats,
            "/switch": self.cmd_switch_tier,
            "/temp": self.cmd_set_temp,
            "/undo": self.cmd_undo,
        }

        func = commands.get(cmd)
        if func:
            return func(args)
        else:
            console.print(f"[bold red]Unknown command:[/bold red] {cmd}")
        return True

    def cmd_branch(self, args):
        new_id = self.session.branch()
        console.print(f"[bold magenta]🌿 Branched to new session: {new_id}[/bold magenta]")
        return True

    def cmd_copy(self, args):
        """Copies the most recent AI response to the system clipboard."""
        import pyperclip
        last_ai = next((m for m in reversed(self.session.history) if m['role'] == 'model'), None)
        if last_ai:
            pyperclip.copy(last_ai['parts'][0])
            console.print("[dim]✅ Copied to clipboard.[/dim]")
        return True

    def cmd_edit(self, args):
        """Allows editing a past prompt with sequential numbering."""
        # 1. Map Display Number -> Real Index
        # This ensures users type '1', '2', '3' regardless of AI turns
        display_to_real = {}
        counter = 1

        if not args:
            table = Table(title="📜 History Timeline (Exchanges)", show_lines=True)
            # CHANGE: Explicitly add column names
            table.add_column("No.", justify="center", style="dim")
            table.add_column("Role", style="bold magenta")
            table.add_column("Content Preview", style="cyan")

            for real_idx, turn in enumerate(self.session.history):
                if turn['role'] == "user":
                    # CHANGE: Ensure we extract the string from the list ['text']
                    parts = turn.get('parts', [""])
                    raw_content = parts[0] if isinstance(parts, list) else parts

                    # Now we can safely call .replace() on the string
                    content = raw_content.replace("\n", " ")

                    if len(content) > 50:
                        content = content[:47] + "..."

                    display_to_real[str(counter)] = real_idx
                    table.add_row(str(counter), "YOU", content)
                    counter += 1

            # Store the map in the router instance so the next call can use it
            self._edit_map = display_to_real

            console.print(table)
            console.print("[dim]Usage: /edit [No.] to travel back to that point.[/dim]")
            return True

        # 2. Pick the index using the map
        display_idx = args[0]
        if hasattr(self, '_edit_map') and display_idx in self._edit_map:
            real_idx = self._edit_map[display_idx]
        else:
            # Fallback to raw index if not in map (or if user knows what they're doing)
            try:
                real_idx = int(display_idx)
            except ValueError:
                console.print("[red]Invalid selection.[/red]")
                return True

        # 3. Capture future prompts before truncating
        # We look for all 'user' turns that happen AFTER the real_idx
        idx = int(real_idx) # Ensure it's an int
        future_prompts = [m for m in self.session.history[idx+1:] if m['role'] == "user"]

        # CHANGE: Extract the string from the list ['text']
        parts = self.session.history[idx]['parts']
        original_text = parts[0] if isinstance(parts, list) else parts

        # 4. Open system editor
        new_text = click.edit(original_text)
        if not new_text or new_text.strip() == original_text.strip():
            console.print("[yellow]No changes saved.[/yellow]")
            return True

        # 5. Strategy Selection
        console.print("\n[bold magenta]🕰️ TIME TRAVEL INITIALIZED[/bold magenta]")
        console.print(" [t] Truncate  [r] Replay  [b] Branch & Truncate")
        while True:
            choice = _pt_session.prompt("Choice [t/r/b] (default t): ").strip().lower() or 't'
            if choice in ('t', 'r', 'b'):
                break
            console.print("[red]Please enter t, r, or b.[/red]")

        if choice == 'b':
            self.session.branch()

        # Perform the actual cut/edit in SessionManager
        self.session.time_travel(real_idx, new_text.strip())

        if choice in ['t', 'b']:
            return {"replay": True, "prompts": []}

        if choice == 'r':
            return {"replay": True, "prompts": future_prompts}


        console.print("✅ History updated.")
        return True

    def cmd_exit(self, args):
        """Terminates the application gracefully."""
        # Optional: Add logic to save current session before closing
        console.print("\n[bold yellow]👋 System offline. Goodbye![/bold yellow]")
        sys.exit(0)

    def cmd_help(self, args):
        """Lists all available commands."""
        table = Table(title="📖 Available Commands", show_lines=True, show_header=True)
        table.add_column("Command", style="bold cyan", no_wrap=True)
        table.add_column("Usage", style="dim")
        table.add_column("Description")
        table.add_column("Status", justify="center")

        all_commands = [
            # (cmd, usage, description, ready)
            ("/search",  "/search [query]",     "Semantic search across all past sessions",                 True),
            ("/resume",  "/resume [No.]",       "List and resume a past session",                           True),
            ("/name",    "/name [name]",        "Set name manually, or omit to auto-name with AI",          True),
            ("/branch",  "",                    "Fork current session into a new branch",                   True),
            ("/edit",    "/edit [No.]",         "Time-travel: edit a past prompt (truncate/replay/branch)", True),
            ("/export",  "/export [last [N]] [--format md|txt|json|pdf|docx] [path]",
                                             "Export last N exchanges in md, txt, json, pdf, or docx",  True),
            ("/history", "/history [No.]",      "Display full conversation of current or a past session",   True),
            ("/undo",    "",                    "Remove the last exchange (You + AI) from history",         True),
            ("/switch",  "/switch [tier]",      "Change model tier mid-session (lite / normal / high)",     True),
            ("/latest",  "",                    "Toggle preference for latest model versions",              True),
            ("/model",   "",                    "Show active model metadata and score details",             True),
            ("/copy",    "",                    "Copy last AI response to clipboard",                       True),
            ("/reset",   "",                    "Clear current session history",                            True),
            ("/stats",   "",                    "Show today\'s token usage across all keys and models",    True),
            ("/save",    "/save [path]",         "Export conversation to a Markdown file",                   True),
            ("/retry",   "",                    "Remove last AI response and resend your prompt",           True),
            ("/new",     "",                    "Save current session and start a fresh one",               False),
            ("/temp",    "/temp [0.0-2.0]",     "Adjust generation temperature",                            False),
            ("/proj",    "/proj [name]",        "Swap the active project prompt profile",                   False),
        ]

        # /help first, then ready sorted, then soon sorted
        ready_cmds = sorted([c for c in all_commands if c[3]], key=lambda x: x[0])
        soon_cmds  = sorted([c for c in all_commands if not c[3]], key=lambda x: x[0])

        table.add_row("/help", "", "Show this help table", "[green]✓[/green]")
        for cmd, usage, desc, ready in ready_cmds + soon_cmds:
            status = "[green]✓[/green]" if ready else "[dim]soon[/dim]"
            table.add_row(cmd, usage, desc, status)

        console.print(table)
        console.print("[dim]Tip: type 'exit' or '/exit' to quit.[/dim]")
        return True

    def cmd_export(self, args):
        """/export [last [N]] [--format md|txt|json] [path]: Export part of the conversation."""
        from pathlib import Path as _Path
        from datetime import datetime as _dt
        import json as _json

        if not self.session.history:
            console.print("[yellow]Nothing to export — session is empty.[/yellow]")
            return True

        # --- Parse arguments ---
        remaining = list(args)
        fmt = "md"
        out_path = None
        n_exchanges = None  # None = all

        # --format flag
        if "--format" in remaining:
            idx = remaining.index("--format")
            if idx + 1 < len(remaining):
                fmt = remaining[idx + 1].lower()
                remaining = remaining[:idx] + remaining[idx + 2:]
            if fmt not in ("md", "txt", "json", "pdf", "docx"):
                console.print("[red]Unknown format. Use: md, txt, or json.[/red]")
                return True

        # 'last' keyword
        if remaining and remaining[0].lower() == "last":
            remaining.pop(0)
            if remaining and remaining[0].isdigit():
                n_exchanges = int(remaining.pop(0))
            else:
                n_exchanges = 1  # default: last 1 exchange

        # Remaining arg = output path
        if remaining:
            out_path = _Path(remaining[0]).expanduser().resolve()

        # --- Slice history ---
        history = self.session.history
        if n_exchanges is not None:
            # Each exchange = 1 user + 1 model turn = 2 entries
            history = history[-(n_exchanges * 2):]

        if not history:
            console.print("[yellow]No exchanges found to export.[/yellow]")
            return True

        # --- Build content ---
        label = f"last_{n_exchanges}" if n_exchanges else "full"
        safe_name = self.session.display_name.replace(" ", "_").replace("/", "-")[:30]

        if fmt == "md":
            ext = ".md"
            lines = [
                f"# {self.session.display_name}",
                f"",
                f"*Exported:* {_dt.now().strftime('%Y-%m-%d %H:%M')}  ",
                f"*Exchanges:* {label}",
                f"",
                "---",
                "",
            ]
            turn_no = 1
            for turn in history:
                role = turn.get("role", "")
                parts = turn.get("parts", [""])
                text = (parts[0] if isinstance(parts, list) else parts).strip()
                if role == "user":
                    lines += [f"## Turn {turn_no} — You", "", text, ""]
                elif role == "model":
                    model_id = turn.get("metadata", {}).get("model", "")
                    label_md = f" *(via {model_id.split('/')[-1]})*" if model_id else ""
                    lines += [f"### AI{label_md}", "", text, ""]
                    turn_no += 1
            content_str = "\n".join(lines)

        elif fmt == "txt":
            ext = ".txt"
            lines = []
            turn_no = 1
            for turn in history:
                role = turn.get("role", "")
                parts = turn.get("parts", [""])
                text = (parts[0] if isinstance(parts, list) else parts).strip()
                if role == "user":
                    lines += [f"[Turn {turn_no}] You:", text, ""]
                elif role == "model":
                    lines += ["AI:", text, ""]
                    turn_no += 1
            content_str = "\n".join(lines)

        elif fmt == "json":
            ext = ".json"
            payload = {
                "session_id": self.session.session_id,
                "display_name": self.session.display_name,
                "exported_at": _dt.now().isoformat(),
                "exchanges": label,
                "history": history,
            }
            content_str = _json.dumps(payload, indent=2, ensure_ascii=False)

        elif fmt == "docx":
            ext = ".docx"
            content_str = None  # docx is written directly to file

        elif fmt == "pdf":
            ext = ".pdf"
            content_str = None  # PDF is written directly to file

        # --- Resolve output path ---
        if out_path is None:
            filename = f"{safe_name}_{label}{ext}"
            out_path = self.session.base_dir / filename
        elif out_path.is_dir():
            out_path = out_path / f"{safe_name}_{label}{ext}"

        out_path.parent.mkdir(parents=True, exist_ok=True)

        if fmt == "docx":
            try:
                from docx import Document as _Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                doc = _Document()

                # Page margins (2cm all sides)
                for section in doc.sections:
                    section.top_margin    = Pt(56)
                    section.bottom_margin = Pt(56)
                    section.left_margin   = Pt(72)
                    section.right_margin  = Pt(72)

                # Title
                title_para = doc.add_heading(self.session.display_name, level=0)
                title_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)

                # Meta line
                meta = doc.add_paragraph()
                meta.add_run(f"Exported: {_dt.now().strftime('%Y-%m-%d %H:%M')}  ·  Exchanges: {label}")
                meta.runs[0].font.size = Pt(9)
                meta.runs[0].font.color.rgb = RGBColor(120, 120, 120)

                doc.add_paragraph()  # spacer

                turn_no = 1
                for turn in history:
                    role  = turn.get("role", "")
                    parts = turn.get("parts", [""])
                    text  = (parts[0] if isinstance(parts, list) else parts).strip()

                    if role == "user":
                        h = doc.add_heading(f"Turn {turn_no} — You", level=2)
                        h.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                        doc.add_paragraph(text)
                    elif role == "model":
                        model_id = turn.get("metadata", {}).get("model", "")
                        model_label = f" (via {model_id.split('/')[-1]})" if model_id else ""
                        h = doc.add_heading(f"AI{model_label}", level=3)
                        h.runs[0].font.color.rgb = RGBColor(46, 125, 50)
                        doc.add_paragraph(text)
                        doc.add_paragraph()  # spacer between exchanges
                        turn_no += 1

                doc.save(str(out_path))
            except ImportError:
                console.print("[red]❌ docx export requires python-docx. Run: uv add python-docx[/red]")
                return True

        elif fmt == "pdf":
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import mm
                from reportlab.lib import colors
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
                from reportlab.lib.enums import TA_LEFT
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from jk_core.constants import THAI_FONT_PATH

                # Register Unicode font if available (supports Thai and other scripts)
                font_name = "Helvetica"  # default fallback
                if THAI_FONT_PATH.exists():
                    try:
                        pdfmetrics.registerFont(TTFont("NotoSansThai", str(THAI_FONT_PATH)))
                        font_name = "NotoSansThai"
                    except Exception:
                        pass  # font broken, fall back to Helvetica

                doc = SimpleDocTemplate(
                    str(out_path),
                    pagesize=A4,
                    leftMargin=20*mm, rightMargin=20*mm,
                    topMargin=20*mm, bottomMargin=20*mm,
                )
                styles = getSampleStyleSheet()
                title_style  = ParagraphStyle("Title",  fontName=font_name, fontSize=18, spaceAfter=6, leading=22)
                h2_style     = ParagraphStyle("H2",     fontName=font_name, fontSize=13, spaceAfter=4, leading=18, spaceBefore=8)
                h3_style     = ParagraphStyle("H3",     fontName=font_name, fontSize=11, spaceAfter=4, leading=16, textColor=colors.HexColor("#2e7d32"))
                body_style   = ParagraphStyle("Body",   fontName=font_name, fontSize=10, spaceAfter=6, leading=16)
                meta_style   = ParagraphStyle("Meta",   fontName=font_name, fontSize=9,  textColor=colors.grey)

                story = [
                    Paragraph(self.session.display_name, title_style),
                    Paragraph(f"Exported: {_dt.now().strftime('%Y-%m-%d %H:%M')} &nbsp;·&nbsp; Exchanges: {label}", meta_style),
                    Spacer(1, 4*mm),
                    HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey),
                    Spacer(1, 4*mm),
                ]

                turn_no = 1
                for turn in history:
                    role  = turn.get("role", "")
                    parts = turn.get("parts", [""])
                    text  = (parts[0] if isinstance(parts, list) else parts).strip()
                    # Escape XML special chars for reportlab
                    safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

                    if role == "user":
                        story.append(Paragraph(f"Turn {turn_no} — You", h2_style))
                        story.append(Paragraph(safe, body_style))
                        story.append(Spacer(1, 2*mm))
                    elif role == "model":
                        model_id = turn.get("metadata", {}).get("model", "")
                        model_label = f" (via {model_id.split('/')[-1]})" if model_id else ""
                        story.append(Paragraph(f"AI{model_label}", h3_style))
                        story.append(Paragraph(safe, body_style))
                        story.append(Spacer(1, 4*mm))
                        turn_no += 1

                doc.build(story)
            except ImportError:
                console.print("[red]❌ PDF export requires reportlab. Run: uv add reportlab[/red]")
                return True
        else:
            out_path.write_text(content_str, encoding="utf-8")

        console.print(f"[green]✅ Exported ({fmt}):[/green] {out_path}")
        return True

    def cmd_history(self, args):
        """/history [No.]: Display the full conversation of a session."""
        import json

        # Resolve which session to show
        if args:
            target_id = args[0]
            if hasattr(self, '_session_map') and target_id in self._session_map:
                target_id = self._session_map[target_id]
            # Load into a temporary dict without clobbering current session
            session_file = self.session.base_dir / f"{target_id}.json"
            if not session_file.exists():
                console.print(f"[red]Session '{target_id}' not found.[/red]")
                console.print("[dim]Tip: run /resume first to build the session list, then use the No.[/dim]")
                return True
            try:
                with open(session_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
            except (json.JSONDecodeError, OSError):
                console.print(f"[red]Session file is corrupted.[/red]")
                return True
            display_name = data.get("display_name", target_id)
            history = data.get("history", [])
        else:
            display_name = self.session.display_name
            history = self.session.history

        if not history:
            console.print("[dim]No messages in this session.[/dim]")
            return True

        console.print(f"\n[bold]📜 {display_name}[/bold] [dim]({len(history)} messages)[/dim]\n")

        turn_no = 1
        for turn in history:
            role = turn.get("role", "")
            parts = turn.get("parts", [""])
            text = parts[0] if isinstance(parts, list) else parts
            text = text.strip()

            if role == "user":
                console.print(f"[bold cyan]  [{turn_no}] You >[/bold cyan] {text}")
            elif role == "model":
                model_id = turn.get("metadata", {}).get("model", "")
                model_label = f"[dim]({model_id.split('/')[-1]})[/dim] " if model_id else ""
                console.print(f"[bold green]      AI >[/bold green] {model_label}{text}")
                turn_no += 1

            console.print()

        return True
    def cmd_model_info(self, args):
        """Shows active model, score breakdown, and full rankings."""
        rankings = self.orch.get_rankings(action="generateContent")
        if not rankings:
            console.print("[yellow]No model rankings available. Run 'jk-ai-init --probe' first.[/yellow]")
            return True

        best = rankings[0]

        # Active model panel
        reasons_str = "\n".join(f"  • {r}" for r in best["reasons"])
        info = (
            f"[bold cyan]Model:[/bold cyan]  [yellow]{best['id']}[/yellow]\n"
            f"[bold cyan]Score:[/bold cyan]  [green]{best['score']}[/green]\n"
            f"[bold cyan]Status:[/bold cyan] {best['status']}\n"
            f"[bold cyan]Tier:[/bold cyan]   {self.orch.tier}\n"
            f"[bold cyan]Logic:[/bold cyan]\n{reasons_str}"
        )
        console.print(Panel(info, title="[bold magenta]Active Model[/bold magenta]", expand=False))

        # Full rankings table
        if len(rankings) > 1:
            table = Table(title="All Ranked Models", show_lines=True)
            table.add_column("Rank", justify="center", style="dim")
            table.add_column("Model", style="cyan")
            table.add_column("Score", justify="right", style="green")
            table.add_column("Status", justify="center")
            for i, m in enumerate(rankings, 1):
                status_color = "green" if m["status"] == "PASS" else "red" if m["status"] == "FAIL" else "dim"
                table.add_row(
                    str(i),
                    m["id"].split("/")[-1],
                    str(m["score"]),
                    f"[{status_color}]{m['status']}[/{status_color}]"
                )
            console.print(table)

        return True

    def cmd_name(self, args):
        """/name [name]: Set name manually, or omit to auto-name using AI."""
        if not args:
            # No args — trigger AI auto-naming
            if not self.session.history:
                console.print("[yellow]Session has no history to name.[/yellow]")
                return True
            return {"autoname": True}

        new_name = " ".join(args)
        self.session.set_display_name(new_name)
        console.print(f"🏷️  Session renamed to: [bold yellow]{new_name}[/bold yellow]")
        return True

    def cmd_new_chat(self, args):
        """
        State Management:
        Saves the current session to a file and starts a fresh one with a new ID.
        """
        console.print("[bold blue]Starting a new chat session...[/bold blue]")
        console.print("[dim]⚠️  /new is not yet implemented.[/dim]")
        return True
    def cmd_reset(self, args):
        self.session.history = []
        self.session.save()
        console.print("[bold green]🧹 Session history cleared.[/bold green]")
        return True

    def cmd_resume(self, args):
        """
        Lists named sessions and loads the selected one.
        Usage: /resume [session_number_or_id]
        """
        recent = self.session.get_recent_sessions(limit=20)
        if not recent:
            console.print("[yellow]No saved sessions found. Use /name first.[/yellow]")
            return True

        # 1. If no ID provided, show the list
        if not args:
            from datetime import datetime
            table = Table(title="🕒 Recently Active Sessions (Top 20)", show_lines=True)
            table.add_column("No.", justify="center", style="dim")
            table.add_column("Name", style="bold yellow")
            table.add_column("Messages", justify="right", style="cyan")
            table.add_column("Last Active", style="dim")

            self._session_map = {}
            for i, (sess_id, info) in enumerate(recent, 1):
                self._session_map[str(i)] = sess_id

                msg_count = str(info.get('message_count', '-'))

                updated_at = info.get('updated_at')
                if updated_at:
                    last_active = datetime.fromtimestamp(updated_at).strftime('%Y-%m-%d %H:%M')
                else:
                    last_active = '-'

                table.add_row(str(i), info['name'], msg_count, last_active)

            console.print(table)
            console.print("[dim]Type '/resume [No.]' to load a session.[/dim]")
            return True

        # 2. Try to load by Number or ID
        target_id = args[0]
        if hasattr(self, '_session_map') and target_id in self._session_map:
            target_id = self._session_map[target_id]

        if self.session.load(target_id):
            console.print(f"✅ [bold green]Resumed session:[/bold green] {self.session.display_name}")
            last_turns = self.session.get_last_turns(n=5)
            if last_turns:
                console.print("\n[bold dim]📜 Recent context (Last 5 exchanges):[/bold dim]")
                for turn in last_turns:
                    role_name = "You" if turn['role'] == "user" else "AI"
                    role_color = "cyan" if turn['role'] == "user" else "green"
                    content_raw = turn['parts'][0].strip().replace("\n", " ")

                    if len(content_raw) > 80:
                        content_raw = content_raw[:77] + "..."

                    line = Text.assemble(
                        (f" {role_name:3}: ", f"bold {role_color}"),
                        (content_raw, "default")
                    )
                    console.print(line)
                console.print("")

            console.print(f"[dim]Total History: {len(self.session.history)} messages loaded.[/dim]")
        else:
            console.print(f"[bold red]Error:[/bold red] Session ID '{target_id}' not found.")
        return True

    def cmd_retry(self, args):
        """Removes the last AI response and re-sends the last user prompt."""
        history = self.session.history
        if len(history) < 2:
            console.print("[yellow]Nothing to retry — no exchange in history.[/yellow]")
            return True

        # Find and remove the last model turn
        if history[-1]["role"] != "model":
            console.print("[yellow]Last message is not from AI — nothing to retry.[/yellow]")
            return True

        self.session.history = history[:-1]
        self.session.save()
        console.print("[dim]↩️  Last AI response removed. Retrying...[/dim]")
        return {"retry": True}

    def cmd_save(self, args):
        """Exports the current conversation to a Markdown file."""
        from pathlib import Path as _Path
        from datetime import datetime as _dt

        if not self.session.history:
            console.print("[yellow]Nothing to save — session is empty.[/yellow]")
            return True

        # Resolve output path
        if args:
            out_path = _Path(args[0]).expanduser().resolve()
        else:
            safe_name = self.session.display_name.replace(" ", "_").replace("/", "-")[:40]
            filename = f"{safe_name}_{self.session.session_id}.md"
            out_path = self.session.base_dir / filename

        # Build Markdown
        lines = [
            f"# {self.session.display_name}",
            f"",
            f"*Session:* `{self.session.session_id}`  ",
            f"*Exported:* {_dt.now().strftime('%Y-%m-%d %H:%M')}",
            f"",
            "---",
            "",
        ]

        turn_no = 1
        for turn in self.session.history:
            role = turn.get("role", "")
            parts = turn.get("parts", [""])
            text = parts[0] if isinstance(parts, list) else parts

            if role == "user":
                lines.append(f"## Turn {turn_no} — You")
                lines.append("")
                lines.append(text.strip())
                lines.append("")
            elif role == "model":
                model_id = turn.get("metadata", {}).get("model", "")
                model_label = f" *(via {model_id.split('/')[-1]})*" if model_id else ""
                lines.append(f"### AI{model_label}")
                lines.append("")
                lines.append(text.strip())
                lines.append("")
                turn_no += 1

        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text("\n".join(lines), encoding="utf-8")
        console.print(f"[green]✅ Saved:[/green] {out_path}")
        return True

    def cmd_search(self, args):
        """/search [query]: Semantic search with pinpoint accuracy."""
        query = " ".join(args)
        if not query.strip():
            console.print("[red]Usage: /search [query][/red]")
            console.print("[dim]Example: /search trip to japan[/dim]")
            console.print("[dim]Tip: searches semantically across all past sessions — typos are fine.[/dim]")
            return True

        se = self.search_engine
        try:
            with console.status("Updating search index..."):
                sessions_updated, turns_added = se.update_index()
            if sessions_updated > 0:
                console.print(f"[dim]🗂  Indexed {turns_added} new turn(s) across {sessions_updated} session(s).[/dim]")
        except Exception as e:
            console.print(f"[bold red]❌ Failed to update index:[/bold red] {e}")
            return True

        results = se.search(query)
        table = Table(title=f"🔍 Semantic Matches for: '{query}'", show_lines=True)
        table.add_column("Score", justify="right", style="green")
        table.add_column("Session / File", style="cyan")
        table.add_column("Matched Turn (The 'Why')", style="yellow")
        for r in results:
            table.add_row(
                f"{r['score']:.2f}",
                f"{r['session_name']}\n[dim]{r['filename']}[/dim]",
                f"Turn {r['turn_no']}: {r['matched_text'][:100]}..."
            )
        if not results:
            console.print("[dim]No matches found.[/dim]")
        console.print(table)
        return True

    def cmd_set_temp(self, args):
        """Adjusts the generation temperature (0.0 to 2.0)."""
        console.print("[dim]⚠️  /temp is not yet implemented.[/dim]")
        return True

    def cmd_stats(self, args):
        """Displays today's accumulated token usage across all keys and models."""
        from datetime import datetime as _dt
        state = self.client.km._load_state()
        usage = state.get("usage", {})

        if not usage:
            console.print("[dim]No usage data recorded yet.[/dim]")
            return True

        table = Table(title="📊 Today's Usage", show_lines=True)
        table.add_column("Key", style="yellow")
        table.add_column("Model", style="cyan")
        table.add_column("Reqs", justify="right", style="magenta")
        table.add_column("In", justify="right", style="green")
        table.add_column("Out", justify="right", style="blue")
        table.add_column("Status", justify="center")
        table.add_column("Since", style="dim")

        total_reqs = total_in = total_out = 0

        for key_id, key_data in sorted(usage.items()):
            models = key_data.get("models", {})
            for model_id, entry in sorted(models.items()):
                reqs = entry.get("request_count", 0)
                inp  = entry.get("total_input_tokens", 0)
                out  = entry.get("total_output_tokens", 0)
                status = entry.get("status", "active")
                window_start = entry.get("window_start", "")

                status_cell = "[green]active[/green]" if status == "active" else "[red]exhausted[/red]"

                since = ""
                if window_start:
                    try:
                        since = _dt.fromisoformat(window_start).astimezone().strftime("%H:%M")
                    except ValueError:
                        pass

                table.add_row(
                    key_id,
                    model_id.split("/")[-1],
                    str(reqs), str(inp), str(out),
                    status_cell, since
                )
                total_reqs += reqs
                total_in   += inp
                total_out  += out

        console.print(table)
        console.print(
            f"[dim]Total today — Reqs: [magenta]{total_reqs}[/magenta]  "
            f"In: [green]{total_in}[/green]  "
            f"Out: [blue]{total_out}[/blue][/dim]"
        )
        return True

    def cmd_switch_proj(self, args):
        """Swaps the Project Profile (Modular Prompts) for the next request."""
        console.print("[dim]⚠️  /proj is not yet implemented.[/dim]")
        return True

    def cmd_switch_tier(self, args):
        """Changes the Orchestrator tier (lite/normal/high) mid-session."""
        if args:
            self.orch.tier = args[0]
            console.print(f"✅ Tier switched to: [bold]{args[0]}[/bold]")
        return True

    def cmd_toggle_latest(self, args):
        """Toggles the 'prefer_latest' flag in the Orchestrator."""
        self.orch.prefer_latest = not self.orch.prefer_latest
        console.print(f"✨ Prefer Latest: [bold]{self.orch.prefer_latest}[/bold]")
        return True

    def cmd_undo(self, args):
        """Removes the last exchange (User + AI) from history."""
        self.session.undo()
        console.print("↩️  Last exchange removed from session.")
        return True

def _assemble_session(proj: str) -> tuple[str, SessionManager]:
    """Returns (system_instruction, session)."""
    variables = {}
    for var in get_required_vars(proj):
        variables[var] = _pt_session.prompt(f"Enter {var.replace('_', ' ').title()}: ")
    system_instruction = assemble_prompt(proj, variables=variables)
    return SessionManager(system_instruction=system_instruction)

def _generate_with_fallback(client, session, rankings):
    """Tries each ranked model in order, returns (resp, meta, used_mid)."""
    for model_entry in rankings:
        target_mid = model_entry['id']
        try:
            with console.status(f"[bold yellow]Thinking ({target_mid.split('/')[-1]})..."):
                resp, meta = client.generate_with_history(
                    history=session.history,
                    system_instruction=session.system_instruction,
                    model_name=target_mid
                )
            return resp, meta, target_mid
        except PermissionError as e:
            if "MODEL_EXHAUSTED" in str(e):
                console.print(f"[dim yellow]⚠️  {target_mid} exhausted. Falling back...[/dim yellow]")
                continue
            raise
    raise RuntimeError("All available models and keys are truly exhausted.")

def _handle_replay(client, session, rankings, future_prompts, model_id):
    """Handles the replay path after a /edit time-travel."""
    try:
        resp, meta, used_mid = _generate_with_fallback(client, session, rankings)
        session.add_message("model", resp, model_id=used_mid)
        console.print(f"[bold green]AI ({client.key_id} | {used_mid.split('/')[-1]}) > [/bold green]{resp}\n")
    except Exception as e:
        console.print(f"[bold red]Replay Error:[/bold red] {e}")
        return

    if not future_prompts:
        return

    console.print(Panel("[bold yellow]🔄 REPLAY MODE ACTIVATED[/bold yellow]\nRegenerating from edited point...", border_style="yellow"))

    for i, p_turn in enumerate(future_prompts, 1):
        p_text = p_turn['parts'][0] if isinstance(p_turn['parts'], list) else p_turn['parts']
        console.print(f"[bold cyan]You (Replay {i}/{len(future_prompts)}) > [/bold cyan]{p_text}")
        session.add_message("user", p_text)
        try:
            resp, meta, used_mid = _generate_with_fallback(client, session, rankings)
            session.add_message("model", resp, model_id=used_mid)
            console.print(f"[bold green]AI ({client.key_id} | {used_mid.split('/')[-1]}) > [/bold green]{resp}\n")
        except Exception as e:
            console.print(f"[bold red]Replay Error:[/bold red] {e}")
            break

    console.print("[bold green]✨ Replay complete. Timeline is now consistent.[/bold green]\n")

def _init_client() -> GeminiClient:
    """Returns an initialized GeminiClient or raises."""
    client = GeminiClient(tier="free")
    client._refresh_client()
    return client

def _print_token_stats(client, meta, model_id):
    """Prints last-request and daily accumulated token usage panels."""
    if not meta:
        return
    from datetime import datetime, timezone
    state = client.km._load_state()
    acc = state.get("usage", {}).get(client.key_id, {}).get("models", {}).get(model_id, {})

    # Label shows which key/model this accumulation is for
    model_short = model_id.split("/")[-1]
    acc_title = f"[bold]Today ({client.key_id} | {model_short})[/bold]"

    # Show window start if available so user knows when the counter began
    window_start = acc.get("window_start")
    window_label = ""
    if window_start:
        try:
            dt = datetime.fromisoformat(window_start).astimezone()
            window_label = f"\n[dim]since {dt.strftime('%H:%M')}[/dim]"
        except ValueError:
            pass

    p1 = Panel(
        f"In:  [green]{meta.prompt_token_count}[/green]\n"
        f"Out: [blue]{meta.candidates_token_count}[/blue]\n"
        f"Tot: [yellow]{meta.total_token_count}[/yellow]",
        title="[bold]Last Request[/bold]", border_style="dim", expand=False
    )
    p2 = Panel(
        f"Reqs: [magenta]{acc.get('request_count', 0)}[/magenta]\n"
        f"In:   [green]{acc.get('total_input_tokens', 0)}[/green]\n"
        f"Out:  [blue]{acc.get('total_output_tokens', 0)}[/blue]"
        f"{window_label}",
        title=acc_title, border_style="dim", expand=False
    )
    console.print(Columns([p1, p2]))

def _do_autoname(client, session):
    """
    Generates and sets a session name using a lite model.
    Uses the full history for context — works at any point in the conversation.
    """
    lite_orch = Orchestrator(provider="google", tier="lite")
    lite_rankings = lite_orch.get_rankings(action="generateContent")
    lite_model = lite_rankings[0]["id"] if lite_rankings else None
    if not lite_model:
        console.print("[yellow]No lite model available for auto-naming.[/yellow]")
        return

    # Build a compact summary of the conversation for the naming prompt
    history = session.history
    first_user = next((m["parts"][0] for m in history if m["role"] == "user"), "")
    last_ai    = next((m["parts"][0] for m in reversed(history) if m["role"] == "model"), "")

    with console.status("[dim]Generating session name...[/dim]"):
        naming_prompt = (
            "Summarize the user's intent in this conversation in "
            "exactly 3 to 5 words. No punctuation. "
            f"\nUser: {first_user[:200]}"
            f"\nAI: {last_ai[:100]}"
        )
        title, _ = client.generate_with_meta(
            prompt=naming_prompt,
            system_instruction="You are a professional filing clerk. Give only the title.",
            model_name=lite_model
        )
        clean_title = title.strip().replace('"', '')
        session.set_display_name(clean_title)
        console.print(
            f"[dim]🏷️  Auto-named via [yellow]{client.key_id}[/yellow] | "
            f"[cyan]{lite_model.split('/')[-1]}[/cyan]: [bold]{clean_title}[/bold][/dim]"
        )


def _run_autoname(client, session, response_text):
    """Auto-names the session after the first exchange (eligibility-gated)."""
    if not session.is_eligible_for_autoname():
        return
    _do_autoname(client, session)
def _read_input(prompt_label: str) -> str:
    """
    Reads user input via prompt_toolkit.

    - Typed input: single line, Enter to submit.
    - Pasted input: bracketed paste mode detects the paste and submits
      the full multi-line block intact automatically.
    - Up/down arrows: navigate input history across turns.
    """
    import re
    from prompt_toolkit.formatted_text import FormattedText
    plain = re.sub(r"\[.*?\]", "", prompt_label)
    # Color only the prefix; leave the input area at terminal default (white)
    styled_prompt = FormattedText([("ansicyan bold", plain)])
    return _pt_session.prompt(styled_prompt, multiline=False)



def _run_chat_loop(client, session, router, rankings, model_id):
    user_input = None

    while True:
        if not user_input:
            prompt_label = f"[bold cyan]You ({client.key_id} | {session.display_name}) > [/bold cyan]"
            user_input = _read_input(prompt_label)

        cmd_result = router.handle(user_input)

        if isinstance(cmd_result, dict) and cmd_result.get("replay"):
            _handle_replay(client, session, rankings, cmd_result["prompts"], model_id)
            user_input = None
            continue

        if isinstance(cmd_result, dict) and cmd_result.get("autoname"):
            try:
                _do_autoname(client, session)
            except Exception as e:
                console.print(f"[red]Auto-name failed:[/red] {e}")
            user_input = None
            continue

        if isinstance(cmd_result, dict) and cmd_result.get("retry"):
            # Re-use the last user message that's now the tail of history
            if session.history and session.history[-1]["role"] == "user":
                parts = session.history[-1].get("parts", [""])
                user_input = parts[0] if isinstance(parts, list) else parts
                # Pop it so the normal path re-adds it cleanly
                session.history = session.history[:-1]
            else:
                user_input = None
            continue

        if cmd_result:
            user_input = None
            continue

        # Normal prompt path
        session.add_message("user", user_input)
        try:
            response_text = ""
            used_mid = model_id

            for candidate in rankings:
                used_mid = candidate['id']
                try:
                    key_before = client.key_id
                    console.print(f"[dim][API] stream_with_history | key={key_before} | model={used_mid.split('/')[-1]}[/dim]")
                    console.print(f"[bold green]AI ({key_before} | {used_mid.split('/')[-1]}) >[/bold green] ", end="")
                    for chunk in client.stream_with_history(
                        history=session.history,
                        system_instruction=session.system_instruction,
                        model_name=used_mid
                    ):
                        print(chunk, end="", flush=True)
                        response_text += chunk
                    print()
                    # If key rotated mid-stream, clarify what actually served it
                    if client.key_id != key_before:
                        console.print(f"[dim](↑ served by {client.key_id} | {used_mid.split('/')[-1]} after key rotation)[/dim]")
                    break  # success
                except PermissionError as e:
                    if "MODEL_EXHAUSTED" in str(e):
                        console.print(f"\n[dim yellow]⚠️  {used_mid} exhausted on all keys. Falling back...[/dim yellow]")
                        response_text = ""
                        continue
                    raise
            else:
                raise RuntimeError("All available models and keys are truly exhausted.")

            session.add_message("model", response_text, model_id=used_mid)
            meta = client.last_meta

            try:
                _run_autoname(client, session, response_text)
            except Exception:
                pass
            _print_token_stats(client, meta, used_mid)
        
        except Exception as e:
            console.print(f"[bold red]Error:[/bold red] {e}")
            session.history.pop()
        
        user_input = None

def _select_model(tier: str, latest: bool):
    """Returns (orchestrator, rankings, model_id) or raises."""
    orch = Orchestrator(provider="google", tier=tier, prefer_latest=latest)
    rankings = orch.get_rankings(action="generateContent")
    if not rankings:
        raise RuntimeError("No healthy models found.")
    best = rankings[0]

    reasons_str = " • " + "\n • ".join(best['reasons'])
    model_info = Text.assemble(
        ("Model: ", "bold cyan"), (f"{best['id']}\n", "yellow"),
        ("Score: ", "bold cyan"), (f"{best['score']}\n", "green"),
        ("Logic:\n", "bold cyan"), (reasons_str, "dim white")
    )
    console.print(Panel(model_info, title="[bold magenta]Selection Logic[/bold magenta]", expand=False))
    return orch, rankings, best['id']

def chat_command(proj, user_input, tier="normal", latest=False):
    try:
        orch, rankings, model_id = _select_model(tier, latest)
        client = _init_client()
        session = _assemble_session(proj)
    except Exception as e:
        console.print(f"[bold red]❌ Error:[/bold red] {e}")
        return

    router = ChatRouter(client, orch, session)
    _run_chat_loop(client, session, router, rankings, model_id)


#def chat_command(proj, user_input, tier="normal", latest=False):
#    orch = Orchestrator(provider="google", tier=tier, prefer_latest=latest)
#    rankings = orch.get_rankings(action="generateContent")
#    if not rankings:
#        console.print("[bold red]❌ Error:[/bold red] No healthy models found.")
#        return
#    best_model = rankings[0]
#    model_id = best_model['id']
#
#    reasons_str = " • " + "\n • ".join(best_model['reasons'])
#    model_info = Text.assemble(
#        (f"Model: ", "bold cyan"), (f"{model_id}\n", "yellow"),
#        (f"Score: ", "bold cyan"), (f"{best_model['score']}\n", "green"),
#        (f"Logic:\n", "bold cyan"), (reasons_str, "dim white")
#    )
#    console.print(Panel(model_info, title="[bold magenta]Selection Logic[/bold magenta]", expand=False))
#
#    # 2. INITIALIZE CLIENT
#    client = GeminiClient(tier="free")
#    try:
#        client._refresh_client()
#    except Exception as e:
#        console.print(f"[bold red]Error:[/bold red] {e}")
#        return
#
#    # 3. PROMPT ASSEMBLY
#    variables = {}
#    required = get_required_vars(proj)
#    for var in required:
#        variables[var] = console.input(f"[bold yellow]Enter {var.replace('_', ' ').title()}: [/bold yellow]")
#    system_instruction = assemble_prompt(proj, variables=variables)
#
#    session = SessionManager(system_instruction=system_instruction)
#    router = ChatRouter(client, orch, session)
#
#    def generate_with_fallback(current_history):
#        # CHANGE: Iterate through all ranked candidates if the top one is exhausted
#        for model_entry in rankings:
#            target_mid = model_entry['id']
#            try:
#                # We show which model is currently being attempted
#                with console.status(f"[bold yellow]Thinking ({target_mid.split('/')[-1]})..."):
#                    return client.generate_with_history(
#                        history=current_history,
#                        system_instruction=session.system_instruction,
#                        model_name=target_mid
#                    ), target_mid
#            except PermissionError as e:
#                if "MODEL_EXHAUSTED" in str(e):
#                    console.print(f"[dim yellow]⚠️  {target_mid} exhausted on all keys. Falling back...[/dim yellow]")
#                    continue
#                raise e
#        raise RuntimeError("All available models and keys are truly exhausted.")
#
#    while True:
#        if not user_input:
#            prompt_label = f"[bold cyan]You ({client.key_id} | {session.display_name}) > [/bold cyan]"
#            user_input = console.input(prompt_label)
#
#        cmd_result = router.handle(user_input)
#        if isinstance(cmd_result, dict) and cmd_result.get("replay"):
#            future_prompts = cmd_result["prompts"]
#            # STEP A: Generate the FIRST response for the newly edited prompt
#            try:
#                # CHANGE: Applied fallback logic to Replay Step A
#                (resp, meta), used_mid = generate_with_fallback(session.history)
#                session.add_message("model", resp, model_id=used_mid)
#                console.print(f"[bold green]AI ({client.key_id} | {used_mid.split('/')[-1]}) > [/bold green]{resp}\n")
#            except Exception as e:
#                console.print(f"[bold red]Replay Error:[/bold red] {e}")
#                user_input = None; continue
#
#            console.print(Panel("[bold yellow]🔄 REPLAY MODE ACTIVATED[/bold yellow]\nRegenerating conversation from the edited point...", border_style="yellow"))
#
#            # STEP B: Replay the REST of the conversation
#            for i, p_turn in enumerate(future_prompts, 1):
#                p_text = p_turn['parts'][0] if isinstance(p_turn['parts'], list) else p_turn['parts']
#                p_model = p_turn.get("metadata", {}).get("model", model_id)
#
#                console.print(f"[bold cyan]You (Replay {i}/{len(future_prompts)}) > [/bold cyan]{p_text}")
#                session.add_message("user", p_text)
#
#                try:
#                    # CHANGE: Applied fallback logic to Replay Step B
#                    (resp, meta), used_mid = generate_with_fallback(session.history)
#                    session.add_message("model", resp, model_id=used_mid)
#                    console.print(f"[bold green]AI ({client.key_id} | {used_mid.split('/')[-1]}) > [/bold green]{resp}\n")
#                except Exception as e:
#                    console.print(f"[bold red]Replay Error:[/bold red] {e}")
#                    break
#
#            console.print("[bold green]✨ Replay complete. Timeline is now consistent.[/bold green]\n")
#            user_input = None
#            continue
#        elif cmd_result:
#            user_input = None
#            continue
#
#        session.add_message("user", user_input)
#        try:
#            console.print(f"\n[bold green]AI ({client.key_id}) > [/bold green]", end="")
#            response_text = ""
#            active_model = model_id
#            for chunk in client.stream_with_history(
#                history=session.history,
#                system_instruction=session.system_instruction,
#                model_name=model_id
#            ):
#                print(chunk, end="", flush=True)
#                response_text += chunk
#            print()  # newline after stream ends
#            session.add_message("model", response_text, model_id=active_model)
#            meta = client.last_meta
#
#            # --- NEW: Milestone 2 - Auto-Naming Logic ---
#            if session.is_eligible_for_autoname():
#                # Use a fast 'lite' model for background tasks to save quota/time
#                with console.status("[dim]Generating session name...[/dim]"):
#                    # We create a specific naming prompt based on the first pair
#                    naming_prompt = (
#                        "Summarize the user's intent in this conversation in "
#                        "exactly 3 to 5 words. No punctuation. "
#                        f"\nUser: {session.history[0]['parts'][0]}"
#                        f"\nAI: {response_text[:100]}"
#                    )
#
#                    # Generate the title
#                    title, _ = client.generate_with_meta(
#                        prompt=naming_prompt,
#                        system_instruction="You are a professional filing clerk. Give only the title.",
#                        model_name=model_id # Or a lite model if you prefer
#                    )
#
#                    clean_title = title.strip().replace('"', '')
#                    session.set_display_name(clean_title)
#                    console.print(f"[dim]🏷️  Session auto-named: [bold]{clean_title}[/bold][/dim]")
#            # --------------------------------------------
#            # 1. Fetch the updated state from disk
#            state = client.km._load_state()
#            acc = state.get("usage", {}).get(client.key_id, {}).get("models", {}).get(model_id, {})
#
#            # 2. Build Last Request Panel
#            last_req_info = (
#                f"In:  [green]{meta.prompt_token_count}[/green]\n"
#                f"Out: [blue]{meta.candidates_token_count}[/blue]\n"
#                f"Tot: [yellow]{meta.total_token_count}[/yellow]"
#            )
#            p1 = Panel(last_req_info, title="[bold]Last Request[/bold]", border_style="dim", expand=False)
#
#            # 3. Build Accumulated Panel (Global)
#            acc_info = (
#                f"Reqs: [magenta]{acc.get('request_count', 0)}[/magenta]\n"
#                f"In:   [green]{acc.get('total_input_tokens', 0)}[/green]\n"
#                f"Out:  [blue]{acc.get('total_output_tokens', 0)}[/blue]"
#            )
#            p2 = Panel(acc_info, title="[bold]Accumulated (Global)[/bold]", border_style="dim", expand=False)
#
#            # 4. Display side-by-side
#            console.print(Columns([p1, p2]))
#            # ----------------------------------------------------
#
#            user_input = None # Clear for next loop iteration
#
#        except Exception as e:
#            console.print(f"[bold red]Error:[/bold red] {e}")
#            session.history.pop()
#            user_input = None

